# -*- coding: utf-8 -*-
"""
Generate Nacravo_SEO_Implementation_Report.docx
Only VERIFIED work is documented. No fabricated rankings/traffic/leads/reviews/
backlinks/CWV/Search Console/GBP/Ads data. Evidence = repository + live production
crawl on 2026-07-21.
"""
from docx import Document
from docx.shared import Pt, Mm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- Brand palette (from site :root vars) ----
MOSS      = RGBColor(0x3B, 0x46, 0x36)
MOSS_DEEP = RGBColor(0x2E, 0x37, 0x2B)
SAGE      = RGBColor(0x7E, 0x8F, 0x70)
SAGE_TEXT = RGBColor(0x5E, 0x6C, 0x4F)
CLAY      = RGBColor(0xB9, 0x87, 0x6B)
PEARL_HEX = "F5F2EC"
LINE_HEX  = "E2DCCF"
MOSS_HEX  = "3B4636"
SAGE_HEX  = "7E8F70"
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
INK       = RGBColor(0x2B, 0x2B, 0x2B)
GREY      = RGBColor(0x6F, 0x6B, 0x63)

DATE = "21 July 2026"

doc = Document()

# ---- Base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for lvl, sz, col in [(1, 17, MOSS), (2, 13.5, MOSS), (3, 11.5, SAGE_TEXT)]:
    st = doc.styles["Heading %d" % lvl]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.bold = True
    st.font.color.rgb = col
    st.paragraph_format.space_before = Pt(12 if lvl == 1 else 8)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.keep_with_next = True

# ---- Page geometry: A4 on the primary section ----
sec = doc.sections[0]
sec.orientation = WD_ORIENT.PORTRAIT
sec.page_width  = Mm(210)
sec.page_height = Mm(297)
sec.top_margin = Mm(20); sec.bottom_margin = Mm(20)
sec.left_margin = Mm(18); sec.right_margin = Mm(18)
sec.different_first_page_header_footer = True  # clean cover page


# ---------- helpers ----------
def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9.5, white=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    r = p.add_run(text if text is not None else "")
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.name = "Calibri"
    if white:
        r.font.color.rgb = WHITE
    elif color is not None:
        r.font.color.rgb = color


def add_table(headers, rows, widths_mm, landscape=False, size=9.0):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    t.autofit = False
    # header
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, MOSS_HEX)
        set_cell_text(c, h, bold=True, white=True, size=size + 0.5)
    # body
    for r_i, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val), size=size)
            if r_i % 2 == 1:
                shade(cells[i], PEARL_HEX)
    # widths (DXA via Mm)
    for i, w in enumerate(widths_mm):
        for row in t.rows:
            row.cells[i].width = Mm(w)
    # spacing after table
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def body(text, bullet=False, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet" if bullet else None)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.font.bold = True
        r.font.color.rgb = MOSS
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def h3(text):
    doc.add_heading(text, level=3)


def field(run_parent, instr):
    """Insert a Word field (e.g. PAGE, NUMPAGES, TOC) into a paragraph."""
    r1 = OxmlElement("w:r"); f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin"); r1.append(f1)
    r2 = OxmlElement("w:r"); it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr; r2.append(it)
    r3 = OxmlElement("w:r"); f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "separate"); r3.append(f3)
    r4 = OxmlElement("w:r"); t4 = OxmlElement("w:t"); t4.text = ""; r4.append(t4)
    r5 = OxmlElement("w:r"); f5 = OxmlElement("w:fldChar"); f5.set(qn("w:fldCharType"), "end"); r5.append(f5)
    for r in (r1, r2, r3, r4, r5):
        run_parent.append(r)


# ---------- running header & footer (applied to primary section) ----------
def build_header_footer(section, title="Nacravo — SEO Implementation Report"):
    hdr = section.header
    hdr.is_linked_to_previous = False
    hp = hdr.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run(title)
    r.font.size = Pt(8); r.font.color.rgb = SAGE_TEXT; r.font.name = "Calibri"
    # bottom border on header paragraph
    pPr = hp._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), LINE_HEX)
    pbdr.append(bottom); pPr.append(pbdr)

    ftr = section.footer
    ftr.is_linked_to_previous = False
    fp = ftr.paragraphs[0]
    fp.text = ""
    # tab stops: left = confidential, right = page
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lr = fp.add_run("Confidential · Nacravo LLC")
    lr.font.size = Pt(8); lr.font.color.rgb = GREY
    fp.add_run("\t\t")
    pr = fp.add_run("Page ")
    pr.font.size = Pt(8); pr.font.color.rgb = GREY
    field(fp._p, "PAGE")
    of = fp.add_run(" of ")
    of.font.size = Pt(8); of.font.color.rgb = GREY
    field(fp._p, "NUMPAGES")


build_header_footer(sec)

# =====================================================================
# COVER PAGE
# =====================================================================
for _ in range(2):
    doc.add_paragraph()

wm = doc.add_paragraph(); wm.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = wm.add_run("nacravo"); r.font.size = Pt(44); r.font.bold = True; r.font.color.rgb = MOSS; r.font.name = "Calibri"

tag = doc.add_paragraph(); tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tag.add_run("Cleaning & Maintenance · Dubai"); r.font.size = Pt(11); r.font.color.rgb = SAGE_TEXT; r.italic = True

# rule
rule = doc.add_paragraph(); rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = rule._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom"); bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), SAGE_HEX); pbdr.append(bottom); pPr.append(pbdr)

for _ in range(2):
    doc.add_paragraph()

title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("SEO Implementation Report"); r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = MOSS_DEEP

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Verified technical, on-page and local SEO work delivered for www.nacravo.com")
r.font.size = Pt(12); r.font.color.rgb = GREY

for _ in range(6):
    doc.add_paragraph()

meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for label, val in [("Prepared for", "Nacravo LLC — Dubai, United Arab Emirates"),
                   ("Production domain", "https://www.nacravo.com"),
                   ("Report date", DATE),
                   ("Evidence basis", "Repository state + live production crawl (21 Jul 2026)")]:
    rr = meta.add_run(label + ":  "); rr.font.bold = True; rr.font.size = Pt(10.5); rr.font.color.rgb = MOSS
    rr2 = meta.add_run(val); rr2.font.size = Pt(10.5); rr2.font.color.rgb = INK
    meta.add_run("\n")

note = doc.add_paragraph(); note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = note.add_run("This report documents only work that was completed and verified. "
                 "No rankings, traffic, leads, reviews, backlinks, Core Web Vitals field data, "
                 "Search Console indexing status, Business Profile changes or Ads performance "
                 "are claimed or fabricated.")
r.font.size = Pt(8.5); r.italic = True; r.font.color.rgb = GREY

doc.add_page_break()

# =====================================================================
# TABLE OF CONTENTS
# =====================================================================
doc.add_heading("Table of Contents", level=1)
tocp = doc.add_paragraph()
field(tocp._p, 'TOC \\o "1-2" \\h \\z \\u')
hint = doc.add_paragraph()
r = hint.add_run("If the entries below appear blank, right-click and choose "
                 "“Update Field” (or press F9) — Word builds the page numbers on open.")
r.font.size = Pt(8.5); r.italic = True; r.font.color.rgb = GREY
doc.add_page_break()

# =====================================================================
# 1. EXECUTIVE SUMMARY
# =====================================================================
h1("1. Executive Summary")
body("Nacravo is a Dubai cleaning and maintenance brand operating a static, Vercel-hosted "
     "website at www.nacravo.com. The objective of the SEO project was to build a crawlable, "
     "conversion-focused site with dedicated service landing pages, correct technical SEO "
     "foundations, and structured data that supports search visibility for Dubai service queries.")

h2("Scope completed (verified)")
body("Eleven service landing pages built, deployed and confirmed live (HTTP 200) in production.", bullet=True)
body("Unique title tags, meta descriptions, canonical tags and H1s on every landing page.", bullet=True)
body("Structured data (Service, HomeAndConstructionBusiness, FAQPage, BreadcrumbList, OfferCatalog) on every landing page; Organization, WebSite, ContactPoint and OpeningHoursSpecification on the homepage.", bullet=True)
body("XML sitemap (23 URLs) and robots.txt published and validated; every sitemap URL returns 200.", bullet=True)
body("Self-referencing canonicals and 15 legacy-path 301 redirects verified in production.", bullet=True)
body("Open Graph / Twitter metadata present; social + schema image corrected to the canonical domain.", bullet=True)
body("Services navigation dropdown defect diagnosed and fixed (accessible, keyboard-operable) and deployed.", bullet=True)

h2("Current production status")
body("All 23 sitemap URLs return HTTP 200 with valid HTTPS on www.nacravo.com. The site is live, "
     "crawlable and indexable. Deployment reconciliation on 21 Jul 2026 brought the previously "
     "un-deployed landing pages into production (they had returned 404 before that date).")

h2("Major outcomes (engineering, not market results)")
body("Eliminated 11 production 404s on the core service/landing pages.", bullet=True)
body("Established a consistent, self-referencing canonical and structured-data layer across the funnel.", bullet=True)
body("Fixed a navigation defect that previously prevented users from reaching service pages from the menu.", bullet=True)

h2("Remaining blockers (require account-owner action)")
body("Apex domain nacravo.com (non-www) fails its TLS certificate check — unresolved (Vercel domain config).", bullet=True)
body("Security headers beyond HSTS are absent (CSP, X-Content-Type-Options, X-Frame-Options, Referrer-Policy, Permissions-Policy).", bullet=True)
body("Search Console indexing/coverage/Core Web Vitals — not verified (authenticated access required).", bullet=True)
body("Google Business Profile — not completed (authenticated access required).", bullet=True)
body("Cross-browser testing on real Safari/Firefox/Edge and live conversion-tracking validation — pending.", bullet=True)

# =====================================================================
# 2. WEBSITE SEO CHANGES IMPLEMENTED
# =====================================================================
doc.add_page_break()
h1("2. Website SEO Changes Implemented")
body("Every row below reflects work present in the repository and confirmed against live production "
     "on 21 Jul 2026. “Verified” describes the method used to confirm it.")

# landscape section for this wide table
ls = doc.add_section(WD_SECTION.NEW_PAGE)
ls.orientation = WD_ORIENT.LANDSCAPE
ls.page_width = Mm(297); ls.page_height = Mm(210)
ls.top_margin = Mm(16); ls.bottom_margin = Mm(16); ls.left_margin = Mm(14); ls.right_margin = Mm(14)
build_header_footer(ls)

rows2 = [
 ["Service landing pages", "11 dedicated service pages created", "Complete", "services, home-cleaning, deep-cleaning, move-in-out-cleaning, holiday-home-cleaning, office-commercial-cleaning, specialized-cleaning, pest-control, ac-service-dubai, handyman-services, annual-maintenance", "Prod crawl: all 200"],
 ["Page titles", "Unique, keyword-led <title> per page", "Complete", "All 11 landing pages + homepage", "Extracted per page (Sec. 3)"],
 ["Meta descriptions", "Unique ~150-160 char descriptions", "Complete", "All 11 landing pages", "Extracted per page (Sec. 3)"],
 ["Canonical tags", "Self-referencing canonical, correct domain", "Complete", "All pages", "9 pages spot-checked live"],
 ["Heading structure", "Single H1 + H2/H3 hierarchy", "Complete", "All landing pages (H2 11-22, H3 14-54)", "Parsed from HTML"],
 ["Internal linking", "Cross-links between services", "Complete", "16 internal service links per page", "Parsed from HTML"],
 ["Sitemap", "XML sitemap, 23 URLs", "Complete", "sitemap.xml", "Well-formed; all 200; no dupes"],
 ["Robots directives", "robots.txt with sitemap ref", "Complete", "robots.txt", "Live 200; /thank-you disallowed"],
 ["Open Graph", "og:title/description/image", "Complete", "index.html + pages", "Live homepage inspected"],
 ["Twitter metadata", "twitter:card summary_large_image", "Complete", "index.html", "Live homepage inspected"],
 ["Structured data", "Service, HomeAndConstructionBusiness, FAQPage, BreadcrumbList, OfferCatalog", "Complete", "All landing pages", "JSON-LD parsed per page"],
 ["Image alt text", "Landing pages use SVG/CSS imagery (0 raster <img>); decorative SVG is aria-hidden", "N/A on LPs", "Landing pages", "0 <img> tags found"],
 ["URL structure", "Clean, keyword-based, cleanUrls", "Complete", "vercel.json cleanUrls:true", "Live URLs verified"],
 ["Redirects", "15 legacy-path 301 redirects", "Complete", "vercel.json", "7 spot-checked: 301 OK"],
 ["Mobile responsiveness", "Responsive nav + layout", "Complete", "index.html + pages", "Breakpoints 360-1920 tested"],
 ["HTTPS", "Valid TLS + HSTS on www", "Complete", "www.nacravo.com", "HSTS max-age=63072000"],
 ["Legal & trust pages", "Privacy, Terms, Cookies, Refund, etc.", "Complete", "9 legal pages", "All live 200"],
 ["Conversion content", "WhatsApp / quote CTAs throughout", "Complete", "5-17 WhatsApp CTAs per page", "Parsed from HTML"],
 ["Navigation fix", "Accessible Services dropdown", "Complete", "index.html", "QA'd + deployed (commit f6f6b7e)"],
 ["Social/schema image", "og/twitter/schema image -> canonical domain", "Fixed", "index.html", "Deployed (commit b61678e)"],
]
add_table(["Area", "Task", "Status", "Files or Pages", "Verification"],
          rows2, widths_mm=[34, 60, 22, 90, 55], size=8.0)

# back to portrait
ps = doc.add_section(WD_SECTION.NEW_PAGE)
ps.orientation = WD_ORIENT.PORTRAIT
ps.page_width = Mm(210); ps.page_height = Mm(297)
ps.top_margin = Mm(20); ps.bottom_margin = Mm(20); ps.left_margin = Mm(18); ps.right_margin = Mm(18)
build_header_footer(ps)

# =====================================================================
# 3. SERVICE LANDING PAGES
# =====================================================================
h1("3. Service Landing Pages")
body("Each page is documented from its actual source and confirmed live. No page is marked live "
     "unless its production URL returned HTTP 200 on 21 Jul 2026 (all did).")

PAGES = [
 ("/services", "All-services hub", "Commercial / navigational", "Cleaning & maintenance services Dubai",
  "home, deep, move-in/out, holiday home, office cleaning; AC, handyman, pest, AMC",
  "Cleaning & Maintenance Services in Dubai | Nacravo",
  "Every Nacravo service in one place — home, deep, move-in/out, holiday home and office cleaning, plus AC, handyman, pest control and annual maintenance in Dubai.",
  "Cleaning and Maintenance Services in Dubai", "WhatsApp quote",
  "16 service links", "Service, OfferCatalog, FAQPage, BreadcrumbList, HomeAndConstructionBusiness", "200"),
 ("/home-cleaning", "Home cleaning", "Commercial (high intent)", "Home cleaning Dubai",
  "maid service, apartment cleaning, trained/insured team",
  "Home Cleaning Services Dubai | Trained, Insured Team",
  "Book home cleaning in Dubai with employed, background-checked technicians. Materials included and a fixed price quoted first. Get your quote on WhatsApp.",
  "Home Cleaning in Dubai", "WhatsApp quote", "16 service links",
  "Service, FAQPage, BreadcrumbList, OfferCatalog, HomeAndConstructionBusiness", "200"),
 ("/deep-cleaning", "Deep cleaning", "Commercial (high intent)", "Deep cleaning Dubai",
  "grease/limescale/grout, apartments & villas, fixed price",
  "Deep Cleaning Services in Dubai | Fixed Price Quote",
  "Deep cleaning in Dubai for apartments and villas. Grease, limescale and grout tackled by insured, employed technicians. Request your fixed price quote today.",
  "Deep Cleaning Services in Dubai", "WhatsApp quote", "16 service links",
  "Service, FAQPage, BreadcrumbList, OfferCatalog, HomeAndConstructionBusiness", "200"),
 ("/move-in-out-cleaning", "Move in / out cleaning", "Commercial (high intent)", "Move in move out cleaning Dubai",
  "end of tenancy, handover standard, landlords/tenants",
  "Move In & Move Out Cleaning Dubai | End of Tenancy",
  "Move in and move out cleaning in Dubai, cleaned to handover standard for landlords and tenants. Photo report supplied. Message us for a fixed price quote.",
  "Move In and Move Out Cleaning in Dubai", "WhatsApp quote", "16 service links",
  "Service, FAQPage, BreadcrumbList, OfferCatalog, HomeAndConstructionBusiness", "200"),
 ("/holiday-home-cleaning", "Holiday home / Airbnb", "Commercial (niche intent)", "Airbnb holiday home cleaning Dubai",
  "turnover cleans, linen change, restocking, hosts",
  "Airbnb & Holiday Home Cleaning Dubai | Book Nacravo",
  "Airbnb and holiday home cleaning in Dubai. Turnover cleans, linen change, restocking and a photo report before check-in. Get a fixed quote on WhatsApp today.",
  "Holiday Home & Airbnb Cleaning in Dubai", "WhatsApp quote", "16 service links",
  "Service, FAQPage, BreadcrumbList, OfferCatalog, HomeAndConstructionBusiness", "200"),
 ("/office-commercial-cleaning", "Office & commercial", "Commercial (B2B)", "Office commercial cleaning Dubai",
  "daily contracts, washrooms, carpets, fit-out handover",
  "Office & Commercial Cleaning Dubai | Nacravo Services",
  "Office and commercial cleaning in Dubai. Daily contracts, deep cleans, washrooms, carpets and fit-out handovers. Request a fixed price quote from Nacravo.",
  "Office & Commercial Cleaning in Dubai", "WhatsApp quote", "16 service links",
  "Service, FAQPage, BreadcrumbList, OfferCatalog, HomeAndConstructionBusiness", "200"),
 ("/specialized-cleaning", "Specialized cleaning", "Commercial (niche intent)", "Sofa carpet mattress cleaning Dubai",
  "sofa, carpet, mattress, curtain, window",
  "Specialized Cleaning Dubai | Sofa, Carpet & Mattress",
  "Specialized cleaning in Dubai for sofas, carpets, mattresses, curtains and interior windows. Machine deep clean with a photo report. Get a fixed quote now.",
  "Specialized Cleaning Services in Dubai", "WhatsApp quote", "16 service links",
  "Service, FAQPage, BreadcrumbList, OfferCatalog, HomeAndConstructionBusiness", "200"),
 ("/pest-control", "Pest control", "Commercial (high intent)", "Pest control Dubai",
  "cockroach, bed bug, ant, rodent treatment",
  "Pest Control Dubai | Cockroach, Bed Bug & Ant Treatment",
  "Pest control in Dubai for cockroaches, bed bugs, ants and rodents. Trained technicians, products used to label instructions. Get a fixed price quote today.",
  "Pest Control Services in Dubai", "WhatsApp quote", "16 service links",
  "Service, FAQPage, BreadcrumbList, OfferCatalog, HomeAndConstructionBusiness", "200"),
 ("/ac-service-dubai", "AC service", "Commercial (high intent, geo)", "AC service Dubai",
  "servicing, chemical wash, duct cleaning, repair; Downtown/Business Bay/DIFC",
  "AC Service Dubai | Servicing, Chemical Wash & Repair",
  "AC servicing, chemical wash, duct cleaning and repair in Downtown Dubai, Business Bay and DIFC. Fixed quote before work, photo report after. Get a quote.",
  "AC Service in Dubai", "WhatsApp quote", "16 service links",
  "Service, FAQPage, BreadcrumbList, OfferCatalog, HomeAndConstructionBusiness", "200"),
 ("/handyman-services", "Handyman", "Commercial (high intent)", "Handyman services Dubai",
  "plumbing, electrical, painting, mounting, door repair",
  "Handyman Services in Dubai | Repairs, Fitting & Fixes",
  "Employed handymen in Dubai for plumbing, electrical, painting, mounting and door repairs. Fixed price quoted before work starts. Book on WhatsApp today.",
  "Handyman Services in Dubai", "WhatsApp quote", "16 service links",
  "Service, FAQPage, BreadcrumbList, OfferCatalog, HomeAndConstructionBusiness", "200"),
 ("/annual-maintenance", "Annual maintenance (AMC)", "Commercial (consideration)", "Annual maintenance contract Dubai",
  "villa/flat AMC, scheduled visits, MEP checks, priority callout",
  "Annual Maintenance Contracts Dubai | Villa & Flat AMC",
  "Annual maintenance contracts in Dubai for apartments and villas: scheduled visits, MEP checks and priority callout for members. Ask for an AMC quote today.",
  "Annual Maintenance Contracts in Dubai", "WhatsApp quote", "16 service links",
  "Service, FAQPage, BreadcrumbList, OfferCatalog, HomeAndConstructionBusiness", "200"),
]

for (url, svc, intent, primary, secondary, title_t, desc, h1t, cta, links, schema, status) in PAGES:
    h2(url)
    rows = [
        ["Target service", svc],
        ["Search intent", intent],
        ["Primary keyword theme", primary],
        ["Secondary keyword themes", secondary],
        ["Title tag", title_t],
        ["Meta description", desc],
        ["H1", h1t],
        ["Main CTA", cta + " (multiple WhatsApp / quote buttons)"],
        ["Internal links", links + " to other Nacravo services + homepage"],
        ["Structured data", schema],
        ["Production HTTP status", status + "  (verified 21 Jul 2026)"],
    ]
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in rows:
        cells = t.add_row().cells
        set_cell_text(cells[0], k, bold=True, color=MOSS, size=9)
        shade(cells[0], PEARL_HEX)
        set_cell_text(cells[1], v, size=9)
    for row in t.rows:
        row.cells[0].width = Mm(42)
        row.cells[1].width = Mm(128)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# =====================================================================
# 4. TECHNICAL SEO AUDIT
# =====================================================================
doc.add_page_break()
h1("4. Technical SEO Audit")
body("Findings from the live production crawl and repository inspection on 21 Jul 2026.")
add_table(["Area", "Finding", "Status"],
    [
     ["Crawlability", "robots.txt allows crawling; nav links are real anchors; sitemap present", "Pass"],
     ["Indexability", "No noindex on money pages; canonical + 200 on all sitemap URLs", "Pass"],
     ["HTTP status codes", "23/23 sitemap URLs return 200; legacy paths 301", "Pass"],
     ["Sitemap integrity", "23 URLs, well-formed XML, no duplicates, all map to live pages", "Pass"],
     ["Robots.txt", "Valid; Sitemap directive present; /thank-you disallowed", "Pass"],
     ["Canonical consistency", "Self-referencing canonicals on canonical www domain", "Pass"],
     ["Redirects", "15 legacy 301s in vercel.json; 7 spot-checked correct", "Pass"],
     ["HTTPS", "Valid TLS on www; HSTS max-age=63072000", "Pass"],
     ["Apex vs www", "nacravo.com (apex) fails TLS cert (SEC_E_WRONG_PRINCIPAL); www valid", "UNRESOLVED"],
     ["Broken links", "No broken internal links found on crawled pages", "Pass"],
     ["Duplicate content", "Unique titles/descriptions/H1s; canonicals prevent duplication", "Pass"],
     ["Security headers", "Only HSTS present; CSP/X-CTO/X-Frame/Referrer/Permissions missing", "UNRESOLVED"],
     ["Mobile rendering", "Responsive; hamburger < 900px; no overflow >= 360px", "Pass"],
     ["Schema validation", "JSON-LD parses; types present per page (manual parse)", "Pass (see note)"],
    ], widths_mm=[38, 108, 24], size=9.0)

h2("Unresolved / limitation notes")
body("Apex-domain certificate: a visitor typing nacravo.com (without www) hits a TLS certificate mismatch. Requires adding the apex domain in Vercel so it provisions a certificate and redirects to www.", bullet=True)
body("Missing security headers: recommend adding a headers block in vercel.json (CSP must be tuned to allow GTM/GA4/Meta before enabling).", bullet=True)
body("Schema validation was performed by parsing JSON-LD from source, not via Google's Rich Results Test (that requires live fetch through Google tooling).", bullet=True)
body("320px viewport: a pre-existing contact-info component (.cinfo) overflows at ~320px; unrelated to navigation. Flagged for a separate responsive fix. (Test harness minimum width was 347px.)", bullet=True)
body("Core Web Vitals field data was not measured (no Lighthouse/PageSpeed run in this session).", bullet=True)

# =====================================================================
# 5. ON-PAGE SEO
# =====================================================================
doc.add_page_break()
h1("5. On-Page SEO")
add_table(["Element", "Implementation (verified)"],
    [
     ["Keyword targeting", "Each landing page targets one primary service theme + related secondaries"],
     ["Search intent alignment", "Commercial-intent pages lead with service, price-first messaging and CTAs"],
     ["Titles", "Unique, keyword-led, brand-suffixed <title> on every page"],
     ["Meta descriptions", "Unique ~150-160 char descriptions with a call to action"],
     ["H1-H3 hierarchy", "Single H1 per page; structured H2/H3 (H2 11-22, H3 14-54)"],
     ["Service-specific content", "Dedicated copy, scope, inclusions and FAQs per service"],
     ["Local Dubai relevance", "Dubai (and area) references in titles, H1s and body copy"],
     ["FAQ content", "FAQPage schema + on-page Q&A on every landing page"],
     ["Trust signals", "Employed/insured team, fixed price, photo report messaging"],
     ["CTA placement", "5-17 WhatsApp / quote CTAs distributed through each page"],
     ["Internal linking", "16 cross-service internal links per page + homepage links"],
     ["Image optimization", "Landing pages use SVG/CSS imagery (no raster <img>); homepage hero is an optimised JPG"],
    ], widths_mm=[46, 124], size=9.5)

# =====================================================================
# 6. LOCAL SEO
# =====================================================================
doc.add_page_break()
h1("6. Local SEO")
h2("Completed (verified)")
add_table(["Area", "Status", "Detail"],
    [
     ["Dubai location targeting", "Complete", "Dubai in titles/H1s/body across pages"],
     ["Service-area content", "Complete", "AC page scoped to Downtown/Business Bay/DIFC; others Dubai-wide"],
     ["NAP consistency", "Complete", "Nacravo LLC, Dubai UAE, +971 58 108 2601, info@nacravo.com sitewide + schema"],
     ["LocalBusiness schema", "Complete", "HomeAndConstructionBusiness + PostalAddress + ContactPoint + OpeningHours"],
     ["Service schema", "Complete", "Service + OfferCatalog on each landing page"],
     ["Local landing pages", "Complete", "11 service pages live"],
    ], widths_mm=[42, 24, 104], size=9.5)
h2("Recommendations / not completed")
body("Google Business Profile: not completed — account access required.", bullet=True, bold_lead="")
body("Review strategy: recommended (post-job WhatsApp review request funnel). Not implemented; no reviews fabricated.", bullet=True)
body("Area-specific landing pages (e.g. per-community) could extend local coverage — recommendation only.", bullet=True)

# =====================================================================
# 7. SEARCH CONSOLE STATUS
# =====================================================================
doc.add_page_break()
h1("7. Search Console Status")
body("All items below require authenticated Google Search Console access, which was not available "
     "in this session. Readiness is reported; verification is not claimed.")
add_table(["Item", "Readiness", "Verified?"],
    [
     ["Sitemap submission", "Ready — sitemap.xml live, valid, all 200", "Not verified — authenticated access required"],
     ["Indexing", "Ready — pages 200, canonical, no noindex", "Not verified — authenticated access required"],
     ["URL Inspection", "Ready to run on homepage + 11 landing pages", "Not verified — authenticated access required"],
     ["Coverage checks", "N/A until data collected", "Not verified — authenticated access required"],
     ["Core Web Vitals", "Not measured this session", "Not verified — authenticated access required"],
     ["Structured data", "JSON-LD present; Rich Results Test pending", "Not verified — authenticated access required"],
    ], widths_mm=[38, 62, 70], size=9.0)

# =====================================================================
# 8. SEO CONTENT STRATEGY
# =====================================================================
doc.add_page_break()
h1("8. SEO Content Strategy")
body("Summary of the content structure already produced on-site. No blog articles are claimed as "
     "published — none exist in production yet; the items below are the live page architecture plus "
     "a recommended publishing order.")
h2("Content pillars (live)")
body("Cleaning services pillar — home, deep, move-in/out, holiday home, office, specialized.", bullet=True)
body("Maintenance services pillar — AC, handyman, annual maintenance.", bullet=True)
body("Pest control pillar.", bullet=True)
h2("Support topics by intent (recommended, not yet published)")
add_table(["Intent", "Example topics", "Status"],
    [
     ["Commercial", "“deep cleaning cost Dubai”, “move-out cleaning checklist”", "Recommended"],
     ["Local", "area guides for Downtown / Business Bay / DIFC AC service", "Recommended"],
     ["Informational", "“how often to service AC in Dubai”, “bed bug signs”", "Recommended"],
    ], widths_mm=[30, 110, 30], size=9.5)
h2("Recommended publishing order")
body("1) High-intent commercial FAQs supporting the top revenue pages (home, deep, AC).", bullet=True)
body("2) Local area content for AC service zones.", bullet=True)
body("3) Informational how-to content to build topical authority.", bullet=True)

# =====================================================================
# 9. BEFORE / AFTER
# =====================================================================
doc.add_page_break()
h1("9. Before-and-After Summary")
add_table(["SEO Area", "Before", "After", "Evidence"],
    [
     ["Service landing pages", "Not in production (404)", "11 pages live (200)", "Prod crawl 21 Jul 2026"],
     ["Sitemap URLs", "11 sitemap URLs 404", "23/23 return 200", "Prod crawl"],
     ["Deployment state", "Local commits unpushed; prod stale", "Merged + deployed to origin/main", "Commits 7219d98, b61678e, f6f6b7e"],
     ["Social/schema image", "Pointed to vercel.app preview domain", "Canonical www.nacravo.com", "Live homepage; commit b61678e"],
     ["Services navigation", "Dropdown closed on hover-gap; links unclickable", "Accessible dropdown; all links clickable", "Browser QA; commit f6f6b7e"],
     ["Canonical domain", "Mixed (preview vs canonical) on image tags", "Consistent canonical domain", "Live inspection"],
    ], widths_mm=[36, 52, 46, 36], size=9.0)

# =====================================================================
# 10. OUTSTANDING TASKS
# =====================================================================
doc.add_page_break()
h1("10. Outstanding Tasks")
add_table(["Priority", "Task", "Reason", "Owner", "Access required"],
    [
     ["Done", "Deploy missing pages", "Pages were 404 in production", "Engineering", "Repo push (completed)"],
     ["Done", "Fix dropdown navigation", "Users could not reach service pages", "Engineering", "Repo push (completed)"],
     ["Done", "Re-crawl production", "Confirm 200s post-deploy", "Engineering", "None (completed)"],
     ["P1", "Resolve apex TLS certificate", "nacravo.com fails cert", "Account owner", "Vercel dashboard"],
     ["P1", "Submit sitemap to Search Console", "Enable indexing", "Account owner", "Search Console login"],
     ["P1", "Inspect URLs in Search Console", "Confirm indexing of 12 pages", "Account owner", "Search Console login"],
     ["P2", "Add security headers", "Hardening (CSP tuned for GTM/GA)", "Engineering", "Repo + config decision"],
     ["P2", "Configure Google Business Profile", "Local pack visibility", "Account owner", "GBP login"],
     ["P2", "Validate live conversion tracking", "Confirm generate_lead / Ads conversion", "Account owner", "GTM Preview + GA4 + Ads"],
     ["P3", "Real Safari & Firefox testing", "Engine-specific QA", "QA", "Physical/BrowserStack devices"],
    ], widths_mm=[18, 46, 44, 28, 34], size=8.5)

# =====================================================================
# 11. FINAL READINESS STATUS
# =====================================================================
doc.add_page_break()
h1("11. Final Readiness Status")
body("Scores are qualitative engineering assessments based on the evidence in this report. They are "
     "NOT measured marketing outcomes (no rankings, traffic, or conversion data is implied).")
add_table(["Dimension", "Score", "Evidence", "Limitations"],
    [
     ["Technical SEO", "8.5 / 10", "23/23 200s; sitemap; canonicals; 301s; HTTPS+HSTS", "Apex cert + security headers unresolved"],
     ["On-page SEO", "9 / 10", "Unique titles/descriptions/H1; schema; internal links", "Rich Results Test not run"],
     ["Local SEO", "6 / 10", "Dubai targeting; LocalBusiness/Service schema; NAP", "GBP not done (account access)"],
     ["Content readiness", "7 / 10", "11 LPs + FAQ architecture live", "No blog/informational content published"],
     ["Measurement readiness", "6 / 10", "GTM/GA4/Ads tags wired in code", "Live tag verification account-gated"],
     ["Production readiness", "9 / 10", "Deployed; all 200; dropdown fixed; no console errors", "Apex cert; real cross-browser pending"],
    ], widths_mm=[40, 20, 66, 44], size=9.0)

h2("Overall")
body("The site is production-ready and technically sound for launch: crawlable, indexable, "
     "structured-data-rich, and conversion-focused. The remaining gaps are (a) two infrastructure "
     "items (apex certificate, security headers) and (b) account-gated activation steps in Search "
     "Console, Business Profile and Ads that require the account owner. None of these block the "
     "site being live and functional today.")

# closing note
doc.add_paragraph()
cl = doc.add_paragraph()
r = cl.add_run("Prepared as an engineering deliverable. All statuses are evidence-backed; account-gated "
               "items are labelled as not verified rather than assumed complete.")
r.font.size = Pt(8.5); r.italic = True; r.font.color.rgb = GREY

import os
out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "Nacravo_SEO_Implementation_Report.docx")
doc.save(out)
print("SAVED:", out)
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables), "sections:", len(doc.sections))
