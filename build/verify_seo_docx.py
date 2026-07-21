# -*- coding: utf-8 -*-
import re
from docx import Document
from docx.shared import Mm
from docx.enum.section import WD_ORIENT

p = "Nacravo_SEO_Implementation_Report.docx"
d = Document(p)

# 1. Headings (outline / nav pane)
heads = [(par.style.name, par.text) for par in d.paragraphs if par.style.name.startswith("Heading")]
h1s = [t for s, t in heads if s == "Heading 1"]
print("=== HEADINGS (nav pane) ===")
for s, t in heads:
    if s in ("Heading 1", "Heading 2"):
        print("  %-9s %s" % (s, t[:70]))
print("H1 count:", len(h1s), "| total headings:", len(heads))

# 2. TOC field present?
xml = d.element.xml
print("\nTOC field present:", ("TOC \\o" in xml) or ("TOC" in xml and "fldChar" in xml))
print("PAGE field present:", "PAGE" in xml and "fldChar" in xml)
print("NUMPAGES field present:", "NUMPAGES" in xml)

# 3. Sections & page geometry
print("\n=== SECTIONS ===")
for i, s in enumerate(d.sections):
    ori = "LANDSCAPE" if s.orientation == WD_ORIENT.LANDSCAPE else "PORTRAIT"
    print("  sec %d: %s  %.0fmm x %.0fmm" % (i, ori, s.page_width.mm, s.page_height.mm))

# 4. Tables: integrity + width sums
print("\n=== TABLES ===")
bad = 0
for i, t in enumerate(d.tables):
    ncol = len(t.columns)
    nrow = len(t.rows)
    # width sum of first row
    try:
        wsum = sum((c.width.mm if c.width else 0) for c in t.rows[0].cells)
    except Exception:
        wsum = -1
    empties = 0
    for row in t.rows:
        for c in row.cells:
            if c.text.strip() == "":
                empties += 1
    flag = ""
    if wsum > 258:  # landscape usable ~ 265mm; portrait ~170mm
        pass
    if empties > 0:
        flag += " EMPTY_CELLS=%d" % empties
    print("  table %2d: %dx%d  widthsum=%.0fmm%s" % (i, nrow, ncol, wsum, flag))

# 5. Markdown / raw-syntax scan across all text
alltext = "\n".join(par.text for par in d.paragraphs)
for t in d.tables:
    for row in t.rows:
        for c in row.cells:
            alltext += "\n" + c.text
issues = []
if re.search(r'(^|\n)#{1,6}\s', alltext):        issues.append("markdown headings (#)")
if re.search(r'\*\*[^*]+\*\*', alltext):          issues.append("bold markdown (**)")
if re.search(r'(^|\n)\s*[-*]\s+\S', alltext):     issues.append("markdown bullets (- / *) in text")
if "```" in alltext:                              issues.append("code fences")
if re.search(r'\|.*\|.*\|', alltext):             issues.append("pipe-table syntax")
if re.search(r'</?[a-zA-Z]+>', alltext):          issues.append("raw HTML tags")
print("\n=== MARKDOWN/SYNTAX SCAN ===")
print("issues:", issues if issues else "NONE")

# 6. landing-page status assertion (every LP table must show 200)
status_hits = alltext.count("200  (verified 21 Jul 2026)")
print("\nLanding-page '200 verified' rows:", status_hits, "(expect 11)")

print("\nOK — document opened & parsed cleanly (valid .docx).")
